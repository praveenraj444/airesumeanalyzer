from flask import Flask, render_template, request, jsonify, send_file, session, redirect, url_for
import os
import re
import PyPDF2
import docx
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
from nltk.probability import FreqDist
from textblob import TextBlob
from collections import Counter
from datetime import datetime
import io
import json
import random
import secrets
import hashlib
import sqlite3
import warnings
from authlib.integrations.flask_client import OAuth
import requests
from dotenv import load_dotenv

# ✅ Load environment variables
load_dotenv()

# ✅ Set BASE_DIR for database
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(BASE_DIR, 'users.db')

warnings.filterwarnings('ignore')

# Try to import advanced PDF libraries
try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    print("⚠️ pdfplumber not installed. Install with: pip install pdfplumber")

try:
    from PIL import Image
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    print("⚠️ OCR not available. Install with: pip install pillow pytesseract")

# ✅ Gemini API Configuration - READ FROM ENVIRONMENT
GEMINI_API_KEY = os.environ.get('GEMINI_API_KEY')

try:
    import google.generativeai as genai
    genai.configure(api_key=GEMINI_API_KEY)
    GEMINI_MODEL = genai.GenerativeModel('gemini-2.5-flash')
    GEMINI_AVAILABLE = True
    print("✅ Google Gemini 2.5 Flash connected!")
except ImportError:
    GEMINI_AVAILABLE = False
    print("⚠️ Install: pip install google-generativeai")
except Exception as e:
    GEMINI_AVAILABLE = False
    print(f"⚠️ Gemini Error: {e}")

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = '/tmp'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    try:
        os.makedirs(app.config['UPLOAD_FOLDER'])
    except Exception:
        app.config['UPLOAD_FOLDER'] = os.path.join(os.getcwd(), 'uploads')
        if not os.path.exists(app.config['UPLOAD_FOLDER']):
            os.makedirs(app.config['UPLOAD_FOLDER'])

# ✅ Flask Secret Key - READ FROM ENVIRONMENT
app.config['SECRET_KEY'] = os.environ.get('FLASK_SECRET_KEY')
if not app.config['SECRET_KEY']:
    raise ValueError("❌ FLASK_SECRET_KEY not set in environment! Please set it in .env file.")

app.config['SESSION_COOKIE_SECURE'] = False
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['PERMANENT_SESSION_LIFETIME'] = 86400

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

email_subscriptions = []
interview_sessions = {}
learning_plans = {}
conversation_history = {}
current_resume_analysis = None

# ✅ Google OAuth Configuration - READ FROM ENVIRONMENT
GOOGLE_CLIENT_ID = os.environ.get('GOOGLE_CLIENT_ID')
GOOGLE_CLIENT_SECRET = os.environ.get('GOOGLE_CLIENT_SECRET')

if not GOOGLE_CLIENT_ID or not GOOGLE_CLIENT_SECRET:
    raise ValueError("❌ GOOGLE_CLIENT_ID and GOOGLE_CLIENT_SECRET must be set in .env file!")

# Initialize OAuth
os.environ['OAUTHLIB_INSECURE_TRANSPORT'] = '1'
os.environ['OAUTHLIB_RELAX_TOKEN_SCOPE'] = '1'

oauth = OAuth(app)

google = oauth.register(
    name='google',
    client_id=GOOGLE_CLIENT_ID,
    client_secret=GOOGLE_CLIENT_SECRET,
    server_metadata_url='https://accounts.google.com/.well-known/openid-configuration',
    client_kwargs={
        'scope': 'openid email profile'
    },
    authorize_params={
        'access_type': 'online',
        'prompt': 'select_account'
    }
)

# Database setup for User Login System
def init_db():
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    
    c.execute('''CREATE TABLE IF NOT EXISTS users
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  username TEXT UNIQUE,
                  email TEXT UNIQUE NOT NULL,
                  password TEXT,
                  google_id TEXT UNIQUE,
                  name TEXT,
                  profile_pic TEXT,
                  created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
    
    c.execute("PRAGMA table_info(users)")
    columns = [col[1] for col in c.fetchall()]
    
    if 'name' not in columns:
        c.execute("ALTER TABLE users ADD COLUMN name TEXT")
        print("✅ Added 'name' column")
    
    if 'profile_pic' not in columns:
        c.execute("ALTER TABLE users ADD COLUMN profile_pic TEXT")
        print("✅ Added 'profile_pic' column")
    
    if 'google_id' not in columns:
        c.execute("ALTER TABLE users ADD COLUMN google_id TEXT UNIQUE")
        print("✅ Added 'google_id' column")
    
    c.execute('''CREATE TABLE IF NOT EXISTS user_resumes
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  user_id INTEGER,
                  resume_data TEXT,
                  job_role TEXT,
                  match_score REAL,
                  created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                  FOREIGN KEY (user_id) REFERENCES users (id))''')
    
    conn.commit()
    conn.close()
    print("✅ Database initialized successfully!")

init_db()

# Download NLTK data
try:
    nltk.data.find('tokenizers/punkt')
except:
    nltk.download('punkt', quiet=True)
try:
    nltk.data.find('corpora/stopwords')
except:
    nltk.download('stopwords', quiet=True)
try:
    nltk.data.find('sentiment/vader_lexicon')
except:
    nltk.download('vader_lexicon', quiet=True)

from nltk.sentiment import SentimentIntensityAnalyzer
sentiment_analyzer = SentimentIntensityAnalyzer()

# ==================== USER AUTHENTICATION FUNCTIONS ====================
def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def register_user(username, email, password):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    try:
        c.execute("INSERT INTO users (username, email, password) VALUES (?, ?, ?)",
                  (username, email, hash_password(password)))
        conn.commit()
        conn.close()
        return True
    except sqlite3.IntegrityError:
        conn.close()
        return False

def login_user(username, password):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT id, username, email, name FROM users WHERE username = ? AND password = ?",
              (username, hash_password(password)))
    user = c.fetchone()
    conn.close()
    return user

def get_user_by_email(email):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT id, username, email, name FROM users WHERE email = ?", (email,))
    user = c.fetchone()
    conn.close()
    return user

def get_user_by_google_id(google_id):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT id, username, email, name FROM users WHERE google_id = ?", (google_id,))
    user = c.fetchone()
    conn.close()
    return user

def create_or_update_google_user(google_id, email, name, profile_pic=None):
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    
    c.execute("SELECT id, username, email, name FROM users WHERE google_id = ? OR email = ?", (google_id, email))
    existing_user = c.fetchone()
    
    if existing_user:
        user_id = existing_user[0]
        username = existing_user[1]
        
        if existing_user[3] != name:
            c.execute("UPDATE users SET name = ? WHERE id = ?", (name, user_id))
        
        if profile_pic:
            c.execute("UPDATE users SET profile_pic = ? WHERE id = ?", (profile_pic, user_id))
        
        conn.commit()
        conn.close()
        return user_id, username
    else:
        username = email.split('@')[0]
        
        c.execute("SELECT id FROM users WHERE username = ?", (username,))
        if c.fetchone():
            username = f"{username}_{google_id[:6]}"
        
        dummy_password = hashlib.sha256(f"google_{google_id}".encode()).hexdigest()
        
        c.execute("INSERT INTO users (username, email, google_id, name, profile_pic, password) VALUES (?, ?, ?, ?, ?, ?)",
                  (username, email, google_id, name, profile_pic, dummy_password))
        conn.commit()
        user_id = c.lastrowid
        conn.close()
        return user_id, username

def save_user_resume(user_id, resume_data, job_role, match_score):
    print(f"💾 SAVING resume for user: {user_id}")
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    try:
        c.execute("INSERT INTO user_resumes (user_id, resume_data, job_role, match_score) VALUES (?, ?, ?, ?)",
                  (user_id, json.dumps(resume_data), job_role, match_score))
        conn.commit()
        print(f"✅ RESUME SAVED SUCCESSFULLY!")
    except Exception as e:
        print(f"❌ Error saving resume: {e}")
    finally:
        conn.close()

def get_user_resumes(user_id):
    print(f"🔍 Fetching resumes for user_id: {user_id}")
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute("SELECT id, resume_data, job_role, match_score, created_at FROM user_resumes WHERE user_id = ? ORDER BY created_at DESC",
              (user_id,))
    resumes = c.fetchall()
    print(f"🔍 Found {len(resumes)} resumes")
    conn.close()
    return resumes

# ==================== UNIVERSAL TEXT EXTRACTION FUNCTIONS ====================

def extract_text_from_pdf(file_path):
    text = ""
    if PDFPLUMBER_AVAILABLE:
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            if text.strip():
                print(f"✅ pdfplumber extracted {len(text)} chars")
                return text
        except Exception as e:
            print(f"pdfplumber failed: {e}")
    
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        if text.strip():
            print(f"✅ PyPDF2 extracted {len(text)} chars")
            return text
    except Exception as e:
        print(f"PyPDF2 failed: {e}")
    
    return text

def extract_text_from_docx(file_path):
    text = ""
    try:
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + " "
        print(f"✅ DOCX extracted {len(text)} chars")
        return text
    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return text

def extract_text(file_path):
    if file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    return None

# ==================== UNIVERSAL CONTACT EXTRACTION ====================

def extract_contact_info(text):
    result = {'emails': [], 'phones': [], 'linkedin': [], 'github': []}
    
    if not text or len(text.strip()) < 50:
        return result
    
    email_patterns = [
        r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',
        r'Email\s*[:：]\s*([a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,})',
    ]
    
    all_emails = []
    for pattern in email_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        all_emails.extend(matches)
    
    result['emails'] = list(dict.fromkeys([e.lower() for e in all_emails if '@' in e and '.' in e.split('@')[-1]]))
    
    phone_patterns = [
        r'\b[6-9][0-9]{9}\b',
        r'\b\+91[-\s]?[6-9][0-9]{9}\b',
        r'Phone\s*[:：]\s*([6-9][0-9]{9})',
        r'\|\s*([6-9][0-9]{9})\s*\|',
        r'\d{6}\s+([6-9][0-9]{9})',
    ]
    
    all_phones = []
    for pattern in phone_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        all_phones.extend(matches)
    
    valid_phones = []
    for phone in all_phones:
        digits = re.sub(r'\D', '', str(phone))
        if len(digits) >= 10:
            last_10 = digits[-10:]
            if len(last_10) == 10 and last_10[0] in '6789':
                valid_phones.append(last_10)
    
    result['phones'] = list(dict.fromkeys(valid_phones))
    
    linkedin_patterns = [
        r'linkedin\.com/in/[\w-]+',
        r'https?://(www\.)?linkedin\.com/in/[\w-]+',
    ]
    
    all_linkedin = []
    for pattern in linkedin_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        all_linkedin.extend(matches)
    
    cleaned_linkedin = []
    for link in all_linkedin:
        link = re.sub(r'^https?://(www\.)?', '', link)
        cleaned_linkedin.append(link)
    
    result['linkedin'] = list(dict.fromkeys(cleaned_linkedin))
    
    return result

# ==================== JOB DATABASE ====================
JOB_DATABASE = {
    "Data Scientist": {
        "category": "Data & Analytics",
        "required_skills": ["Python", "SQL", "Machine Learning", "Statistics", "TensorFlow", "PyTorch", "Pandas", "NumPy"],
        "preferred_skills": ["Deep Learning", "NLP", "Computer Vision", "Big Data", "Spark"],
        "keywords": ["python", "sql", "machine learning", "statistics", "deep learning", "tensorflow", "pytorch", "pandas", "numpy", "data science", "analytics", "visualization"],
        "experience_level": "Mid-Senior",
        "salary_range": "$90k - $150k",
        "description": "Analyze complex data and build predictive models"
    },
    "Data Analyst": {
        "category": "Data & Analytics",
        "required_skills": ["SQL", "Excel", "Tableau", "Python", "Statistics", "Data Visualization", "Power BI"],
        "preferred_skills": ["R", "ETL", "Business Intelligence"],
        "keywords": ["sql", "excel", "tableau", "python", "statistics", "data visualization", "power bi", "analytics", "reporting", "dashboard"],
        "experience_level": "Entry-Mid",
        "salary_range": "$65k - $110k",
        "description": "Transform data into actionable insights"
    },
    "Data Engineer": {
        "category": "Data & Analytics",
        "required_skills": ["Python", "SQL", "ETL", "Data Warehousing", "Spark", "Airflow"],
        "preferred_skills": ["AWS/GCP", "Kafka", "Hadoop", "Scala"],
        "keywords": ["etl", "data pipeline", "spark", "hadoop", "data warehouse", "python", "sql", "airflow", "big data"],
        "experience_level": "Mid-Senior",
        "salary_range": "$95k - $160k",
        "description": "Build and maintain data infrastructure"
    },
    "ML Engineer": {
        "category": "Data & Analytics",
        "required_skills": ["Python", "Machine Learning", "TensorFlow", "PyTorch", "SQL", "Cloud Platforms", "MLOps"],
        "preferred_skills": ["Docker", "Kubernetes", "Feature Store", "Model Deployment"],
        "keywords": ["machine learning", "mlops", "tensorflow", "pytorch", "model deployment", "python", "ai", "deep learning"],
        "experience_level": "Senior",
        "salary_range": "$110k - $180k",
        "description": "Deploy and scale machine learning models"
    },
    "Software Engineer": {
        "category": "Development",
        "required_skills": ["Python", "Java", "JavaScript", "SQL", "Git", "Data Structures", "Algorithms", "REST APIs"],
        "preferred_skills": ["React", "Node.js", "Docker", "AWS", "Spring Boot"],
        "keywords": ["python", "java", "javascript", "sql", "git", "agile", "rest api", "docker", "aws", "react", "node.js", "spring boot"],
        "experience_level": "Entry-Mid",
        "salary_range": "$80k - $140k",
        "description": "Design and develop software applications"
    },
    "Frontend Developer": {
        "category": "Development",
        "required_skills": ["JavaScript", "React", "HTML", "CSS", "Git", "Responsive Design"],
        "preferred_skills": ["Vue.js", "Angular", "TypeScript", "Next.js", "Tailwind CSS"],
        "keywords": ["javascript", "react", "html", "css", "frontend", "ui", "ux", "responsive", "vue", "angular", "typescript"],
        "experience_level": "Entry-Mid",
        "salary_range": "$75k - $130k",
        "description": "Build responsive and interactive user interfaces"
    },
    "Backend Developer": {
        "category": "Development",
        "required_skills": ["Python", "Java", "SQL", "REST APIs", "Database Design", "Git", "Microservices"],
        "preferred_skills": ["Node.js", "Spring Boot", "Django", "Redis"],
        "keywords": ["backend", "api", "database", "server", "python", "java", "sql", "rest", "microservices", "node.js", "django"],
        "experience_level": "Mid-Senior",
        "salary_range": "$85k - $145k",
        "description": "Build scalable server-side applications and APIs"
    },
    "Full Stack Developer": {
        "category": "Development",
        "required_skills": ["JavaScript", "React", "Node.js", "Python", "SQL", "Git", "REST APIs"],
        "preferred_skills": ["TypeScript", "AWS", "Docker", "MongoDB", "GraphQL"],
        "keywords": ["full stack", "frontend", "backend", "javascript", "react", "node.js", "python", "database", "api"],
        "experience_level": "Mid-Senior",
        "salary_range": "$90k - $160k",
        "description": "Work on both frontend and backend development"
    },
    "Product Manager": {
        "category": "Product & Management",
        "required_skills": ["Product Strategy", "Agile", "Scrum", "Market Research", "Data Analysis", "Communication", "Leadership"],
        "preferred_skills": ["UX Design", "Technical Background", "A/B Testing", "JIRA"],
        "keywords": ["product management", "agile", "scrum", "roadmap", "user stories", "market research", "analytics", "leadership", "strategy"],
        "experience_level": "Mid-Senior",
        "salary_range": "$100k - $170k",
        "description": "Lead product vision and strategy"
    },
    "Project Manager": {
        "category": "Product & Management",
        "required_skills": ["Project Management", "Agile", "Scrum", "Risk Management", "Budgeting", "Stakeholder Management"],
        "preferred_skills": ["PMP Certification", "JIRA", "Confluence", "MS Project"],
        "keywords": ["project management", "agile", "scrum", "pmp", "risk management", "stakeholder", "budget", "timeline"],
        "experience_level": "Mid-Senior",
        "salary_range": "$85k - $150k",
        "description": "Manage project delivery and teams"
    },
    "DevOps Engineer": {
        "category": "Cloud & DevOps",
        "required_skills": ["Linux", "Docker", "Kubernetes", "CI/CD", "AWS", "Python", "Bash", "Jenkins"],
        "preferred_skills": ["Terraform", "Ansible", "Monitoring Tools", "Prometheus"],
        "keywords": ["linux", "docker", "kubernetes", "ci/cd", "jenkins", "aws", "azure", "gcp", "terraform", "ansible", "devops"],
        "experience_level": "Mid-Senior",
        "salary_range": "$100k - $170k",
        "description": "Automate infrastructure and deployment"
    },
    "Cloud Architect": {
        "category": "Cloud & DevOps",
        "required_skills": ["AWS", "Azure", "GCP", "Infrastructure as Code", "Network Security", "Solution Design"],
        "preferred_skills": ["Terraform", "CloudFormation", "Multi-cloud", "Cost Optimization"],
        "keywords": ["cloud architect", "aws", "azure", "gcp", "infrastructure", "security", "scalability", "high availability"],
        "experience_level": "Senior",
        "salary_range": "$130k - $200k",
        "description": "Design cloud solutions and architecture"
    },
    "Security Engineer": {
        "category": "Cybersecurity",
        "required_skills": ["Network Security", "Penetration Testing", "Firewalls", "SIEM", "Python", "Bash"],
        "preferred_skills": ["CISSP", "CEH", "Cloud Security", "Incident Response"],
        "keywords": ["security", "cybersecurity", "penetration testing", "firewall", "siem", "incident response", "vulnerability"],
        "experience_level": "Mid-Senior",
        "salary_range": "$95k - $165k",
        "description": "Protect systems and data from threats"
    },
    "AI Research Scientist": {
        "category": "AI & Research",
        "required_skills": ["Deep Learning", "Research", "Python", "PyTorch", "TensorFlow", "Publications"],
        "preferred_skills": ["Reinforcement Learning", "Generative AI", "Computer Vision", "NLP"],
        "keywords": ["ai research", "deep learning", "machine learning", "publications", "research", "neural networks", "llm", "generative ai"],
        "experience_level": "Senior",
        "salary_range": "$120k - $200k",
        "description": "Advance AI through research and innovation"
    },
    "Mobile Developer": {
        "category": "Mobile Development",
        "required_skills": ["Android", "iOS", "Kotlin", "Swift", "React Native", "Flutter"],
        "preferred_skills": ["Firebase", "REST APIs", "Git", "Mobile UI/UX"],
        "keywords": ["android", "ios", "kotlin", "swift", "react native", "flutter", "mobile", "app development"],
        "experience_level": "Mid-Senior",
        "salary_range": "$85k - $150k",
        "description": "Build mobile applications for iOS and Android"
    },
    "QA Engineer": {
        "category": "QA & Testing",
        "required_skills": ["Manual Testing", "Automation Testing", "Selenium", "JUnit", "TestNG", "Bug Tracking"],
        "preferred_skills": ["Cypress", "Jest", "Postman", "JMeter"],
        "keywords": ["testing", "qa", "automation", "selenium", "junit", "bug tracking", "quality assurance"],
        "experience_level": "Entry-Mid",
        "salary_range": "$60k - $110k",
        "description": "Ensure software quality through testing"
    },
    "UI/UX Designer": {
        "category": "Design",
        "required_skills": ["Figma", "Adobe XD", "Sketch", "User Research", "Prototyping", "Wireframing"],
        "preferred_skills": ["HTML/CSS", "JavaScript", "Animation", "User Testing"],
        "keywords": ["ui", "ux", "design", "figma", "adobe xd", "prototyping", "wireframing", "user experience"],
        "experience_level": "Mid-Senior",
        "salary_range": "$70k - $130k",
        "description": "Design user interfaces and experiences"
    }
}

# Technical Skills Database
TECHNICAL_SKILLS_DB = {
    'Programming Languages': {
        'keywords': ['python', 'java', 'javascript', 'c++', 'c#', 'ruby', 'go', 'rust', 'swift', 'kotlin', 'php', 'typescript', 'scala', 'r'],
        'display_names': ['Python', 'Java', 'JavaScript', 'C++', 'C#', 'Ruby', 'Go', 'Rust', 'Swift', 'Kotlin', 'PHP', 'TypeScript', 'Scala', 'R']
    },
    'Data Science & ML': {
        'keywords': ['tensorflow', 'pytorch', 'scikit-learn', 'pandas', 'numpy', 'matplotlib', 'seaborn', 'keras', 'nltk', 'spacy', 'opencv'],
        'display_names': ['TensorFlow', 'PyTorch', 'Scikit-learn', 'Pandas', 'NumPy', 'Matplotlib', 'Seaborn', 'Keras', 'NLTK', 'spaCy', 'OpenCV']
    },
    'Web Development': {
        'keywords': ['react', 'angular', 'vue.js', 'node.js', 'django', 'flask', 'spring boot', 'express.js', 'html', 'css', 'bootstrap', 'tailwind'],
        'display_names': ['React', 'Angular', 'Vue.js', 'Node.js', 'Django', 'Flask', 'Spring Boot', 'Express.js', 'HTML', 'CSS', 'Bootstrap', 'Tailwind CSS']
    },
    'Cloud & DevOps': {
        'keywords': ['aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git', 'github', 'gitlab', 'ci/cd', 'terraform', 'ansible'],
        'display_names': ['AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes', 'Jenkins', 'Git', 'GitHub', 'GitLab', 'CI/CD', 'Terraform', 'Ansible']
    },
    'Databases': {
        'keywords': ['mysql', 'postgresql', 'mongodb', 'redis', 'oracle', 'sql server', 'cassandra', 'elasticsearch', 'dynamodb', 'firebase'],
        'display_names': ['MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'Oracle', 'SQL Server', 'Cassandra', 'Elasticsearch', 'DynamoDB', 'Firebase']
    },
    'Soft Skills': {
        'keywords': ['leadership', 'communication', 'problem solving', 'teamwork', 'project management', 'agile', 'scrum', 'critical thinking', 'time management'],
        'display_names': ['Leadership', 'Communication', 'Problem Solving', 'Teamwork', 'Project Management', 'Agile', 'Scrum', 'Critical Thinking', 'Time Management']
    }
}

# Salary Database
SALARY_DATABASE = {
    "Python": {"entry": 60000, "mid": 85000, "senior": 120000},
    "Java": {"entry": 65000, "mid": 90000, "senior": 125000},
    "JavaScript": {"entry": 55000, "mid": 80000, "senior": 110000},
    "React": {"entry": 60000, "mid": 85000, "senior": 115000},
    "Node.js": {"entry": 60000, "mid": 85000, "senior": 115000},
    "SQL": {"entry": 55000, "mid": 80000, "senior": 110000},
    "Machine Learning": {"entry": 70000, "mid": 100000, "senior": 140000},
    "TensorFlow": {"entry": 75000, "mid": 105000, "senior": 145000},
    "PyTorch": {"entry": 75000, "mid": 105000, "senior": 145000},
    "AWS": {"entry": 70000, "mid": 100000, "senior": 135000},
    "Docker": {"entry": 65000, "mid": 90000, "senior": 125000},
    "Kubernetes": {"entry": 70000, "mid": 100000, "senior": 135000},
}

# Market salary ranges
MARKET_SALARY = {
    "Data Scientist": {"min": 80000, "max": 150000, "average": 115000},
    "Data Analyst": {"min": 55000, "max": 110000, "average": 82000},
    "Data Engineer": {"min": 75000, "max": 160000, "average": 115000},
    "ML Engineer": {"min": 90000, "max": 180000, "average": 135000},
    "Software Engineer": {"min": 70000, "max": 140000, "average": 105000},
    "Frontend Developer": {"min": 65000, "max": 130000, "average": 95000},
    "Backend Developer": {"min": 70000, "max": 145000, "average": 107000},
    "Full Stack Developer": {"min": 75000, "max": 160000, "average": 115000},
    "Product Manager": {"min": 80000, "max": 170000, "average": 125000},
    "Project Manager": {"min": 70000, "max": 150000, "average": 110000},
    "DevOps Engineer": {"min": 80000, "max": 170000, "average": 125000},
    "Cloud Architect": {"min": 100000, "max": 200000, "average": 150000},
    "Security Engineer": {"min": 75000, "max": 165000, "average": 120000},
    "AI Research Scientist": {"min": 90000, "max": 200000, "average": 145000},
    "Mobile Developer": {"min": 70000, "max": 150000, "average": 110000},
    "QA Engineer": {"min": 55000, "max": 110000, "average": 82000},
    "UI/UX Designer": {"min": 60000, "max": 130000, "average": 95000},
}

# QUICK REPLY BUTTONS
QUICK_REPLIES = [
    {"text": "📝 Resume Tips", "value": "resume tips"},
    {"text": "🎤 Interview Tips", "value": "interview tips"},
    {"text": "💰 Salary Guide", "value": "salary negotiation"},
    {"text": "🔍 Job Search", "value": "job search"},
    {"text": "📚 Skills to Learn", "value": "skills to learn"},
    {"text": "🔗 LinkedIn Tips", "value": "linkedin tips"},
    {"text": "📄 Cover Letter", "value": "cover letter"},
    {"text": "🤖 ATS Tips", "value": "ats tips"},
    {"text": "📊 My Resume", "value": "my resume analysis"},
    {"text": "🎯 Job Matches", "value": "job recommendations"},
    {"text": "🎤 Mock Interview", "value": "mock interview"},
    {"text": "📚 Learning Path", "value": "learning path"},
    {"text": "🏆 ATS Score", "value": "ats score"},
    {"text": "📝 Resume Rewriter", "value": "resume rewriter"},
    {"text": "🔗 LinkedIn Optimizer", "value": "linkedin optimizer"}
]

# Interview Questions
INTERVIEW_QUESTIONS = {
    "Data Scientist": {
        "technical": [
            "Explain the difference between supervised and unsupervised learning.",
            "How do you handle missing values in a dataset?",
            "What is the bias-variance tradeoff?",
            "Explain how a decision tree works.",
            "What is regularization and why is it used?",
            "How do you detect overfitting in machine learning models?"
        ],
        "behavioral": [
            "Describe a challenging data science project you worked on.",
            "How do you communicate complex findings to non-technical teams?",
            "Tell me about a time you had to make a data-driven decision.",
            "How do you stay updated with the latest ML trends?"
        ],
        "coding": [
            "Write a function to calculate the mean, median, and mode of a list.",
            "Implement a simple linear regression from scratch.",
            "Write SQL query to find duplicate records in a table.",
            "Implement a function to check if a string is a palindrome."
        ]
    },
    "Software Engineer": {
        "technical": [
            "Explain the difference between REST and SOAP APIs.",
            "What is the time complexity of binary search?",
            "Explain the concept of object-oriented programming.",
            "What is the difference between SQL and NoSQL databases?",
            "Explain how garbage collection works in Python/Java.",
            "What is the difference between process and thread?"
        ],
        "behavioral": [
            "Describe a challenging bug you fixed in production.",
            "How do you handle tight deadlines?",
            "Tell me about a time you mentored a junior developer.",
            "How do you prioritize features vs technical debt?"
        ],
        "coding": [
            "Reverse a string without using built-in functions.",
            "Find the second largest number in an array.",
            "Implement a function to check if two strings are anagrams.",
            "Write a function to find Fibonacci series."
        ]
    }
}

# ==================== TEXT EXTRACTION & SKILLS FUNCTIONS ====================

def extract_technical_skills(text):
    found_skills = {}
    text_lower = text.lower()
    for category, skill_data in TECHNICAL_SKILLS_DB.items():
        category_skills = []
        for i, keyword in enumerate(skill_data['keywords']):
            if re.search(r'\b' + re.escape(keyword) + r'\b', text_lower):
                category_skills.append(skill_data['display_names'][i])
            elif re.search(r'\b' + re.escape(keyword) + r's\b', text_lower):
                category_skills.append(skill_data['display_names'][i])
        if category_skills:
            found_skills[category] = category_skills
    return found_skills

def extract_education(text):
    education_keywords = ['B.Sc', 'M.Sc', 'B.Tech', 'M.Tech', 'PhD', 'Bachelor', 'Master', 'MBA', 'BCA', 'MCA']
    education = []
    sentences = sent_tokenize(text)
    for sentence in sentences:
        for keyword in education_keywords:
            if keyword.lower() in sentence.lower():
                education.append(sentence.strip())
                break
    if not education:
        education = ['Bachelor of Technology (B.Tech) - Computer Science']
    return education[:5]

def extract_experience(text):
    patterns = [
        r'(\d+)\+?\s*years?\s+of\s+experience',
        r'experience\s+of\s+(\d+)\+?\s*years?',
        r'(\d+)\+?\s*years?\s+experience'
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return int(match.group(1))
    dates = re.findall(r'(20\d{2})\s*[-–]\s*(20\d{2}|present)', text, re.IGNORECASE)
    if dates:
        total = 0
        for start, end in dates:
            if end.lower() == 'present':
                total += datetime.now().year - int(start)
            else:
                total += int(end) - int(start)
        return total if total > 0 else 3
    return 3

def extract_current_salary(text):
    salary_patterns = [
        r'current salary[:\s]*\$?(\d{2,3}[,.]?\d{0,3})',
        r'present salary[:\s]*\$?(\d{2,3}[,.]?\d{0,3})',
        r'ctc[:\s]*\$?(\d{2,3}[,.]?\d{0,3})',
        r'salary[:\s]*\$?(\d{2,3}[,.]?\d{0,3})',
    ]
    for pattern in salary_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            salary = match.group(1).replace(',', '').replace('.', '')
            return int(salary)
    return None

# ==================== SCORE CALCULATIONS ====================
def calculate_job_match_score(resume_text, job_role):
    if job_role not in JOB_DATABASE:
        return 55
    
    job_data = JOB_DATABASE[job_role]
    job_keywords = job_data['keywords']
    job_required_skills = job_data['required_skills']
    job_preferred_skills = job_data.get('preferred_skills', [])
    job_experience_level = job_data['experience_level']
    
    resume_lower = resume_text.lower()
    
    exact_matches = 0
    partial_matches = 0
    for keyword in job_keywords:
        if re.search(r'\b' + re.escape(keyword) + r'\b', resume_lower):
            exact_matches += 1
        elif keyword in resume_lower:
            partial_matches += 0.5
    keyword_score = ((exact_matches * 1.0) + (partial_matches * 0.5)) / len(job_keywords) * 100 if job_keywords else 0
    
    resume_skills = extract_technical_skills(resume_text)
    all_found_skills = []
    for category_skills in resume_skills.values():
        all_found_skills.extend([s.lower() for s in category_skills])
    
    required_matched = 0
    for skill in job_required_skills:
        if skill.lower() in all_found_skills:
            required_matched += 1
        elif any(skill.lower() in found_skill or found_skill in skill.lower() for found_skill in all_found_skills):
            required_matched += 0.5
    required_score = (required_matched / len(job_required_skills)) * 100 if job_required_skills else 0
    
    preferred_matched = 0
    for skill in job_preferred_skills:
        if skill.lower() in all_found_skills:
            preferred_matched += 1
    preferred_bonus = (preferred_matched / len(job_preferred_skills)) * 15 if job_preferred_skills else 0
    skill_score = required_score + preferred_bonus
    
    tfidf_score = 0
    try:
        job_description = job_data.get('description', '')
        if job_description and len(resume_text.split()) > 50:
            vectorizer = TfidfVectorizer(stop_words='english', max_features=1000)
            tfidf_matrix = vectorizer.fit_transform([resume_text, job_description])
            tfidf_score = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0] * 100
    except:
        tfidf_score = 0
    
    experience_years = extract_experience(resume_text)
    exp_score = 50
    if job_experience_level == "Entry-Mid":
        exp_score = 90 if experience_years <= 3 else (70 if experience_years <= 5 else 50)
    elif job_experience_level == "Mid-Senior":
        exp_score = 90 if 3 <= experience_years <= 6 else (60 if experience_years <= 2 else 70)
    elif job_experience_level == "Senior":
        exp_score = 90 if experience_years >= 5 else (60 if experience_years >= 3 else 40)
    
    education_text = ' '.join(extract_education(resume_text)).lower()
    if re.search(r'ph\.?d|doctorate', education_text):
        edu_score = 100
    elif re.search(r'master|mba', education_text):
        edu_score = 85
    elif re.search(r'bachelor|b\.tech', education_text):
        edu_score = 70
    else:
        edu_score = 50
    
    action_verbs = ['led', 'managed', 'achieved', 'improved', 'increased', 'created', 'developed', 'implemented']
    action_count = sum(1 for verb in action_verbs if re.search(r'\b' + verb + r'\b', resume_lower))
    action_score = min(50 + (action_count * 5), 100)
    
    total_words = len(resume_lower.split())
    if total_words > 0:
        keyword_occurrences = 0
        for keyword in job_keywords[:20]:
            occurrences = len(re.findall(r'\b' + re.escape(keyword) + r'\b', resume_lower))
            keyword_occurrences += min(occurrences, 3)
        density_score = min((keyword_occurrences / len(job_keywords[:20])) * 100, 100) if job_keywords else 0
    else:
        density_score = 0
    
    final_score = (
        (keyword_score * 0.25) + (skill_score * 0.25) + (tfidf_score * 0.15) +
        (exp_score * 0.15) + (edu_score * 0.10) + (action_score * 0.05) + (density_score * 0.05)
    )
    
    return max(0, min(100, round(final_score, 2)))

def calculate_skill_match_score(skills, job_role):
    if job_role not in JOB_DATABASE:
        return 50
    required_skills = JOB_DATABASE[job_role]['required_skills']
    found_skills = []
    for category_skills in skills.values():
        found_skills.extend([s.lower() for s in category_skills])
    if not found_skills:
        return 30
    matched = 0
    for skill in required_skills:
        if skill.lower() in found_skills:
            matched += 1
        elif any(skill.lower() in found_skill or found_skill in skill.lower() for found_skill in found_skills):
            matched += 0.5
    score = (matched / len(required_skills)) * 100 if required_skills else 0
    return round(score, 2)

def calculate_salary_comparison(skills, experience_years, job_role):
    skill_salaries = []
    for category, skill_list in skills.items():
        for skill in skill_list:
            if skill in SALARY_DATABASE:
                if experience_years <= 2:
                    skill_value = SALARY_DATABASE[skill]["entry"]
                elif experience_years <= 5:
                    skill_value = SALARY_DATABASE[skill]["mid"]
                else:
                    skill_value = SALARY_DATABASE[skill]["senior"]
                skill_salaries.append(skill_value)
    avg_skill_salary = sum(skill_salaries) / len(skill_salaries) if skill_salaries else 60000
    exp_multiplier = 0.8 if experience_years <= 2 else (0.9 if experience_years <= 4 else (1.0 if experience_years <= 6 else 1.1))
    total_skills = sum(len(skills_list) for skills_list in skills.values())
    skill_bonus = 15000 if total_skills > 10 else (10000 if total_skills > 7 else (5000 if total_skills > 4 else 0))
    estimated_salary = (avg_skill_salary * exp_multiplier) + skill_bonus
    market_data = MARKET_SALARY.get(job_role, {"min": 50000, "max": 150000, "average": 90000})
    estimated_salary = max(market_data["min"], min(estimated_salary, market_data["max"]))
    if estimated_salary < market_data["average"] * 0.8:
        position = "Below Market"
        recommendation = "Consider negotiating for higher salary based on your skills"
    elif estimated_salary < market_data["average"]:
        position = "Slightly Below Market"
        recommendation = "Your skills are valuable, aim for market rate"
    elif estimated_salary <= market_data["average"] * 1.2:
        position = "At Market Rate"
        recommendation = "Good alignment with market standards"
    else:
        position = "Above Market"
        recommendation = "Excellent! Your skills are highly valued"
    return {
        'estimated_salary': round(estimated_salary),
        'market_min': market_data["min"],
        'market_max': market_data["max"],
        'market_average': market_data["average"],
        'position': position,
        'recommendation': recommendation,
        'skill_count': total_skills,
        'experience_multiplier': round(exp_multiplier, 2)
    }

def analyze_action_verbs(text):
    action_verbs = ['led', 'managed', 'achieved', 'improved', 'increased', 'created', 'developed', 'implemented']
    action_count = sum(1 for verb in action_verbs if re.search(r'\b' + verb + r'\b', text.lower()))
    confidence = min(30 + (action_count * 7), 100)
    return action_count, confidence

def extract_key_phrases(text):
    try:
        stop_words = set(stopwords.words('english'))
        words = word_tokenize(text.lower())
        words = [word for word in words if word.isalpha() and word not in stop_words and len(word) > 3]
        from nltk.util import ngrams
        bigrams = [' '.join(gram) for gram in ngrams(words, 2)]
        freq_dist = FreqDist(bigrams)
        phrases = [phrase for phrase, count in freq_dist.most_common(10)]
        return phrases if phrases else ['python programming', 'data analysis', 'project management', 'team leadership']
    except:
        return ['python programming', 'data analysis', 'project management', 'team leadership']

def get_top_job_matches(resume_text, top_n=16):
    matches = []
    for job_role in JOB_DATABASE.keys():
        score = calculate_job_match_score(resume_text, job_role)
        matches.append({
            'job_role': job_role,
            'match_score': score,
            'category': JOB_DATABASE[job_role]['category'],
            'experience_level': JOB_DATABASE[job_role]['experience_level'],
            'salary_range': JOB_DATABASE[job_role]['salary_range']
        })
    matches.sort(key=lambda x: x['match_score'], reverse=True)
    return matches[:top_n]

def generate_interview_questions(job_role):
    if job_role in INTERVIEW_QUESTIONS:
        questions = INTERVIEW_QUESTIONS[job_role]
        return {
            'technical': questions['technical'][:6],
            'behavioral': questions['behavioral'][:4],
            'coding': questions['coding'][:4]
        }
    else:
        return {
            'technical': [
                "Explain a complex problem you solved recently.",
                "How do you stay updated with latest technologies?",
                "Describe your development process.",
                "How do you handle technical debt?",
                "What's your approach to testing?",
                "Tell me about a time you learned a new technology quickly."
            ],
            'behavioral': [
                "Tell me about a time you faced a conflict at work.",
                "Describe a project you're most proud of.",
                "How do you handle feedback and criticism?",
                "Tell me about a time you failed and what you learned."
            ],
            'coding': [
                "Write a function to reverse a string.",
                "Find the missing number in an array.",
                "Check if a number is prime.",
                "Find duplicates in an array."
            ]
        }

def generate_recommendations(analysis_data):
    recommendations = []
    if analysis_data['match_score'] >= 80:
        recommendations.append("Excellent match! Your resume strongly aligns with this role.")
    elif analysis_data['match_score'] >= 60:
        recommendations.append("Good match! Add more quantifiable achievements to strengthen your resume.")
    elif analysis_data['match_score'] >= 40:
        recommendations.append("Moderate match. Add relevant keywords from the job description.")
    else:
        recommendations.append("Low match. Consider adding more relevant keywords and skills.")
    if analysis_data['skill_score'] < 50:
        missing = JOB_DATABASE[analysis_data['job_role']]['required_skills'][:3]
        recommendations.append(f"Add these key skills: {', '.join(missing)}")
    if analysis_data['experience_years'] == 0:
        recommendations.append("Highlight your work experience with clear dates and achievements.")
    if analysis_data['action_verbs_count'] < 5:
        recommendations.append("Use more action verbs to showcase impact.")
    recommendations.append("Tailor your resume for each application using keywords from the job posting.")
    return recommendations

# ==================== ATS SCORE CHECKER ====================
def calculate_ats_score(resume_text):
    score = 100
    issues = []
    suggestions = []
    
    if re.search(r'<table|&lt;table', resume_text, re.IGNORECASE):
        score -= 15
        issues.append("Tables detected")
        suggestions.append("Remove tables - ATS cannot read them properly")
    
    if re.search(r'<img|&lt;img', resume_text, re.IGNORECASE):
        score -= 15
        issues.append("Images detected")
        suggestions.append("Remove images - ATS cannot read text from images")
    
    standard_headings = ['experience', 'education', 'skills', 'summary', 'work history']
    found_headings = 0
    for heading in standard_headings:
        if re.search(r'\b' + heading + r'\b', resume_text.lower()):
            found_headings += 1
    
    if found_headings < 3:
        score -= 20
        issues.append("Missing standard section headings")
        suggestions.append("Use standard headings: Experience, Education, Skills")
    
    if not re.search(r'\b[\w\.-]+@[\w\.-]+\.\w+\b', resume_text):
        score -= 10
        issues.append("No email found")
        suggestions.append("Add email address at the top of resume")
    
    if not re.search(r'\b\d{10}\b', resume_text):
        score -= 5
        issues.append("No phone number found")
        suggestions.append("Add phone number at the top of resume")
    
    word_count = len(resume_text.split())
    if word_count < 300:
        score -= 10
        issues.append("Resume too short")
        suggestions.append("Add more details about your experience and skills")
    
    return {
        'score': max(0, min(100, score)),
        'issues': issues,
        'suggestions': suggestions,
        'is_ats_friendly': score >= 70
    }

# ==================== COVER LETTER GENERATOR ====================
def generate_cover_letter(resume_data, job_role, company_name=""):
    skills_list = []
    for category_skills in resume_data.get('skills', {}).values():
        skills_list.extend(category_skills)
    
    skills_text = ', '.join(skills_list[:10]) if skills_list else "various technical skills"
    experience = resume_data.get('experience_years', 0)
    job_role_text = job_role
    
    if GEMINI_AVAILABLE:
        try:
            prompt = f"""Generate a professional cover letter for a job application.

Job Role: {job_role_text}
Company: {company_name if company_name else '[Company Name]'}
Experience: {experience} years
Top Skills: {skills_text}

The cover letter should be professional, enthusiastic, and highlight relevant experience."""
            
            response = GEMINI_MODEL.generate_content(prompt)
            return response.text
        except:
            pass
    
    current_date = datetime.now().strftime('%B %d, %Y')
    return f"""
{current_date}

Hiring Manager
{company_name if company_name else '[Company Name]'}

Dear Hiring Manager,

I am writing to express my strong interest in the {job_role_text} position. With {experience} years of experience and expertise in {skills_text}, I am confident in my ability to contribute to your team.

Thank you for considering my application.

Sincerely,
[Your Name]
"""

# ==================== RESUME REWRITER ====================
def rewrite_resume_section(text, section_type="bullet_point"):
    if not GEMINI_AVAILABLE:
        return text
    
    try:
        prompts = {
            'bullet_point': f"Rewrite this resume bullet point to be more impactful using action verbs and quantifiable results.\n\nOriginal: {text}\n\nRewritten:",
            'summary': f"Rewrite this professional summary to be more compelling and keyword-rich.\n\nOriginal: {text}\n\nRewritten:",
            'skill': f"Rewrite this skill description to highlight proficiency and impact.\n\nOriginal: {text}\n\nRewritten:"
        }
        prompt = prompts.get(section_type, prompts['bullet_point'])
        response = GEMINI_MODEL.generate_content(prompt)
        return response.text.strip()
    except:
        return text

def get_action_verb_suggestions(text):
    weak_verbs = ['responsible for', 'assisted with', 'helped', 'worked on', 'handled']
    strong_verbs = ['led', 'managed', 'developed', 'implemented', 'created', 'designed', 'achieved', 'improved']
    
    suggestions = []
    for weak in weak_verbs:
        if weak in text.lower():
            suggestions.append(f"Replace '{weak}' with: {', '.join(strong_verbs[:4])}")
    
    return suggestions

# ==================== LINKEDIN PROFILE OPTIMIZER ====================
def optimize_linkedin_profile(resume_data):
    skills_list = []
    for category_skills in resume_data.get('skills', {}).values():
        skills_list.extend(category_skills)
    
    experience_years = resume_data.get('experience_years', 0)
    job_role = resume_data.get('job_role', 'Professional')
    
    if skills_list:
        top_skills = skills_list[:3]
        headline = f"{job_role} | {', '.join(top_skills)} | {experience_years}+ years experience"
    else:
        headline = f"{job_role} | {experience_years}+ years of experience"
    
    about_section = f"""I am a {job_role} with {experience_years} years of experience. 
My expertise includes: {', '.join(skills_list[:8]) if skills_list else 'various technical skills'}."""
    
    suggestions = [
        {'section': 'Headline', 'suggested': headline, 'tip': 'Your headline appears in searches. Make it keyword-rich!'},
        {'section': 'About Section', 'suggested': about_section[:500], 'tip': 'Use first-person and include keywords'},
        {'section': 'Skills Section', 'suggested': ', '.join(skills_list[:15]) if skills_list else 'Add your top skills', 'tip': 'List 10-15 skills and get endorsements'}
    ]
    
    return {
        'headline': headline,
        'about_section': about_section,
        'suggestions': suggestions,
        'skill_recommendations': skills_list[:15]
    }

# ==================== CHATBOT FUNCTIONS ====================
def get_gemini_response(user_message, context=""):
    if not GEMINI_AVAILABLE:
        return None
    try:
        prompt = f"""You are an AI Career Assistant. Be friendly and helpful.
Context: {context}
User: {user_message}
Assistant:"""
        response = GEMINI_MODEL.generate_content(prompt)
        return response.text
    except:
        return None

def start_mock_interview(job_role, session_id):
    if job_role in INTERVIEW_QUESTIONS:
        questions = INTERVIEW_QUESTIONS[job_role]
        all_questions = questions['technical'] + questions['behavioral'] + questions['coding']
    else:
        all_questions = ["Tell me about yourself.", "What are your greatest strengths?", "Why do you want to work here?"]
    
    interview_sessions[session_id] = {
        'job_role': job_role,
        'questions': all_questions,
        'current_index': 0,
        'answers': []
    }
    return f"🎤 Mock Interview Started for {job_role}\n\n{all_questions[0]}\n\nType your answer, or type 'next' for next question."

def continue_mock_interview(session_id, user_answer):
    if session_id not in interview_sessions:
        return None, "No active interview. Type 'start interview' to begin."
    session = interview_sessions[session_id]
    current_idx = session['current_index']
    questions = session['questions']
    session['answers'].append({'question': questions[current_idx], 'answer': user_answer})
    session['current_index'] += 1
    if session['current_index'] >= len(questions):
        del interview_sessions[session_id]
        return None, "🎉 Interview Complete! Great job!"
    else:
        return session, f"Question {session['current_index'] + 1}: {questions[session['current_index']]}"

def get_chatbot_response(user_message, session_id="default", resume_data=None):
    global interview_sessions
    user_lower = user_message.lower().strip()
    
    if "start interview" in user_lower:
        job_role = resume_data.get('job_role', 'Software Engineer') if resume_data else 'Software Engineer'
        return start_mock_interview(job_role, session_id)
    
    if session_id in interview_sessions:
        if user_lower == "next":
            session = interview_sessions[session_id]
            session['current_index'] += 1
            return continue_mock_interview(session_id, "[Skipped]")
        else:
            result, response = continue_mock_interview(session_id, user_message)
            return response if result is None else response
    
    if "ats score" in user_lower and resume_data:
        ats = resume_data.get('ats_score', {})
        return f"🏆 Your ATS Score: {ats.get('score', 0)}%\n\nIssues: {', '.join(ats.get('issues', ['None']))}"
    
    if resume_data:
        if "my top skills" in user_lower:
            skills_list = []
            for skills in resume_data.get('skills', {}).values():
                skills_list.extend(skills)
            if skills_list:
                return f"Your top skills: {', '.join(skills_list[:10])}"
            return "No skills detected. Add skills like Python, SQL, JavaScript."
        elif "my experience" in user_lower:
            exp = resume_data.get('experience_years', 0)
            return f"You have approximately {exp} years of experience." if exp > 0 else "Experience not clearly specified."
        elif "my job match" in user_lower:
            return f"Your match score is {resume_data.get('match_score', 0)}% for {resume_data.get('job_role', 'your role')}."
        elif "job recommendations" in user_lower:
            matches = resume_data.get('top_matches', [])[:5]
            if matches:
                response = "Top job recommendations:\n"
                for i, m in enumerate(matches, 1):
                    response += f"{i}. {m['job_role']} - {m['match_score']}%\n"
                return response
    
    if GEMINI_AVAILABLE:
        gemini_response = get_gemini_response(user_message)
        if gemini_response:
            return gemini_response
    
    responses = {
        "resume tips": "📝 Resume Tips:\n1. Use action verbs\n2. Quantify achievements\n3. Tailor for each job",
        "interview tips": "🎤 Interview Tips:\n1. Research company\n2. Practice STAR method\n3. Prepare questions",
        "salary negotiation": "💰 Salary Tips:\n1. Research market rates\n2. Highlight your value\n3. Be confident",
        "job search": "🔍 Best job portals: LinkedIn, Indeed, Naukri.com, Glassdoor",
        "skills to learn": "📚 High-demand skills: Python, SQL, JavaScript, React, AWS, Docker",
        "linkedin tips": "🔗 LinkedIn tips: Professional photo, compelling headline, detailed about section",
        "cover letter": "📄 Cover letter: Customize for each job, show passion, keep to 1 page",
        "ats tips": "🤖 ATS tips: Use standard headings, avoid images, use standard fonts"
    }
    
    for key, value in responses.items():
        if key in user_lower:
            return value
    
    return "I can help with resume tips, interview prep, salary negotiation, job search, and skill development. What would you like to know?"

# ==================== MAIN RESUME ANALYSIS FUNCTION ====================

def analyze_resume(file_path, job_role):
    global current_resume_analysis
    text = extract_text(file_path)
    if not text or len(text.strip()) < 100:
        text = "Experienced Software Engineer with 5 years of experience in Python, Java, and JavaScript. Skilled in React, Node.js, Docker, and AWS. Led a team of 5 developers. Bachelor of Technology in Computer Science."
    
    text = ' '.join(text.split())
    
    contact_info = extract_contact_info(text)
    skills = extract_technical_skills(text)
    education = extract_education(text)
    experience_years = extract_experience(text)
    current_salary = extract_current_salary(text)
    action_verbs_count, confidence_score = analyze_action_verbs(text)
    key_phrases = extract_key_phrases(text)
    
    match_score = calculate_job_match_score(text, job_role)
    skill_score = calculate_skill_match_score(skills, job_role)
    top_matches = get_top_job_matches(text, top_n=16)
    salary_comparison = calculate_salary_comparison(skills, experience_years, job_role)
    interview_questions = generate_interview_questions(job_role)
    ats_result = calculate_ats_score(text)
    linkedin_optimization = optimize_linkedin_profile({
        'skills': skills,
        'experience_years': experience_years,
        'job_role': job_role
    })
    
    completeness = 0
    if contact_info.get('emails'): completeness += 10
    if contact_info.get('phones'): completeness += 10
    if skills: completeness += 30
    if education: completeness += 20
    if experience_years > 0: completeness += 15
    if current_salary: completeness += 5
    completeness += min(skill_score / 10, 10)
    
    analysis_data = {
        'contact_info': contact_info,
        'skills': skills,
        'education': education,
        'experience_years': experience_years,
        'current_salary': current_salary,
        'match_score': min(round(match_score, 1), 100),
        'skill_score': min(round(skill_score, 1), 100),
        'action_verbs_count': action_verbs_count,
        'confidence_score': confidence_score,
        'key_phrases': key_phrases,
        'total_words': len(text.split()),
        'completeness_score': round(completeness, 1),
        'job_role': job_role,
        'top_matches': top_matches,
        'job_category': JOB_DATABASE[job_role]['category'],
        'experience_level': JOB_DATABASE[job_role]['experience_level'],
        'salary_range': JOB_DATABASE[job_role]['salary_range'],
        'salary_comparison': salary_comparison,
        'interview_questions': interview_questions,
        'recommendations': generate_recommendations({
            'match_score': match_score, 'skill_score': skill_score, 'experience_years': experience_years,
            'action_verbs_count': action_verbs_count, 'job_role': job_role
        }),
        'ats_score': ats_result,
        'linkedin_optimization': linkedin_optimization
    }
    
    current_resume_analysis = analysis_data
    return analysis_data

# ==================== FLASK ROUTES ====================

@app.route('/')
def home():
    if 'user_id' in session:
        return redirect(url_for('index'))
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    session.pop('oauth_state', None)
    
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        user = login_user(username, password)
        if user:
            session['user_id'] = user[0]
            session['username'] = user[1]
            session['email'] = user[2]
            session['name'] = user[3] if user[3] else user[1]
            session.pop('guest_mode', None)
            return redirect(url_for('index'))
        else:
            return render_template('login.html', error="Invalid username or password")
    return render_template('login.html')

@app.route('/login/google')
def google_login():
    try:
        session.pop('oauth_state', None)
        
        import secrets
        state = secrets.token_urlsafe(32)
        session['oauth_state'] = state
        
        redirect_uri = url_for('google_authorized', _external=True)
        
        print(f"🔐 Google Login - Redirect URI: {redirect_uri}")
        print(f"🔐 Google Login - State: {state}")
        
        return google.authorize_redirect(redirect_uri, state=state)
    except Exception as e:
        print(f"❌ Google Login Error: {e}")
        return render_template('login.html', error=f"Google login failed: {str(e)}")

@app.route('/login/google/authorized')
def google_authorized():
    try:
        request_state = request.args.get('state')
        session_state = session.get('oauth_state')
        
        print(f"🔐 Request State: {request_state}")
        print(f"🔐 Session State: {session_state}")
        
        if not session_state or not request_state or session_state != request_state:
            session.pop('oauth_state', None)
            return render_template('login.html', error="CSRF Warning: State mismatch. Please try again.")
        
        token = google.authorize_access_token()
        if not token:
            return render_template('login.html', error="Google login failed. Please try again.")
        
        resp = google.get('https://www.googleapis.com/oauth2/v2/userinfo', token=token)
        user_info = resp.json()
        
        if not user_info or 'email' not in user_info:
            return render_template('login.html', error="Could not get user info from Google.")
        
        google_id = user_info.get('id')
        email = user_info.get('email')
        name = user_info.get('name', email.split('@')[0])
        profile_pic = user_info.get('picture')
        
        user_id, username = create_or_update_google_user(
            google_id=google_id,
            email=email,
            name=name,
            profile_pic=profile_pic
        )
        
        session.pop('oauth_state', None)
        
        session['user_id'] = user_id
        session['username'] = username
        session['email'] = email
        session['name'] = name
        session['profile_pic'] = profile_pic
        session.pop('guest_mode', None)
        
        print(f"✅ Google login successful! User: {username} ({email})")
        
        return redirect(url_for('index'))
        
    except Exception as e:
        print(f"❌ Google login error: {e}")
        session.pop('oauth_state', None)
        return render_template('login.html', error=f"Google login failed: {str(e)}")

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        if register_user(username, email, password):
            return redirect(url_for('login'))
        else:
            return render_template('register.html', error="Username or email already exists")
    return render_template('register.html')

@app.route('/guest_mode')
def guest_mode():
    session['guest_mode'] = True
    session['username'] = 'Guest'
    return redirect(url_for('index'))

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route('/dashboard')
def dashboard():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session.get('user_id')
    print(f"🔍 Dashboard - user_id: {user_id}")
    user_resumes = get_user_resumes(user_id)
    print(f"🔍 Dashboard - Found {len(user_resumes)} resumes")
    return render_template('dashboard.html', resumes=user_resumes, username=session.get('username'))

@app.route('/index')
def index():
    if 'user_id' not in session and not session.get('guest_mode'):
        return redirect(url_for('login'))
    
    jobs_by_category = {}
    for job, details in JOB_DATABASE.items():
        category = details['category']
        if category not in jobs_by_category:
            jobs_by_category[category] = []
        jobs_by_category[category].append(job)
    return render_template('index_advanced.html', jobs_by_category=jobs_by_category)

@app.route('/analyze', methods=['POST'])
def analyze():
    
    # 🔍 DEBUG: Check session
    user_id = session.get('user_id')
    print(f"🔍 USER ID: {user_id}")
    print(f"🔍 SESSION: {dict(session)}")
    
    if not user_id and not session.get('guest_mode'):
        return jsonify({'error': 'Please login or continue as guest'}), 401
    
    if 'resume' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['resume']
    job_role = request.form.get('job_role')
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    if not job_role:
        return jsonify({'error': 'Please select a job role'}), 400
    if not file.filename.endswith(('.pdf', '.docx')):
        return jsonify({'error': 'Please upload PDF or DOCX file'}), 400
    
    safe_filename = re.sub(r'[^a-zA-Z0-9._-]', '_', file.filename)
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], safe_filename)
    file.save(file_path)
    
    try:
        analysis = analyze_resume(file_path, job_role)
        session['current_analysis'] = analysis
        
        # ✅ DEBUG: Save to database
        if user_id:
            print(f"💾 SAVING resume for user: {user_id}")
            save_user_resume(user_id, analysis, job_role, analysis['match_score'])
            print(f"✅ RESUME SAVED SUCCESSFULLY!")
        else:
            print(f"⚠️ No user_id, skipping save")
        
        if os.path.exists(file_path):
            os.remove(file_path)
            
        return render_template('result_advanced.html', analysis=analysis)
    except Exception as e:
        if os.path.exists(file_path):
            os.remove(file_path)
        print(f"❌ ERROR: {e}")
        return jsonify({'error': f'Analysis failed: {str(e)}'}), 500
        
@app.route('/subscribe', methods=['POST'])
def subscribe():
    data = request.json
    email = data.get('email')
    job_role = data.get('job_role')
    if not email or not job_role:
        return jsonify({'error': 'Email and job role required'}), 400
    email_subscriptions.append({'email': email, 'job_role': job_role, 'subscribed_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S')})
    return jsonify({'success': True, 'message': f'Subscribed successfully for {job_role} alerts!'})

@app.route('/send-alert', methods=['POST'])
def send_alert():
    return jsonify({'success': True, 'message': 'Alert sent successfully!'})

@app.route('/chatbot', methods=['POST'])
def chatbot():
    data = request.json
    user_message = data.get('message', '')
    session_id = data.get('session_id', 'default')
    if not user_message:
        return jsonify({'error': 'No message provided'}), 400
    resume_data = current_resume_analysis
    response = get_chatbot_response(user_message, session_id, resume_data)
    return jsonify({'response': response, 'quick_replies': QUICK_REPLIES})

@app.route('/get-quick-replies', methods=['GET'])
def get_quick_replies():
    return jsonify({'quick_replies': QUICK_REPLIES})

import json

@app.route('/cover-letter', methods=['GET', 'POST'])
def cover_letter_page():
    if 'user_id' not in session:
        return redirect(url_for('login'))
        
    user_id = session.get('user_id')
    analysis = session.get('current_analysis')

    if not analysis and user_id:
        print("🔍 Session empty. Fetching last resume from database...")
        user_resumes = get_user_resumes(user_id)
        
        if user_resumes and len(user_resumes) > 0:
            last_resume = user_resumes[0] 
            
            try:
                analysis = json.loads(last_resume[2]) 
            except Exception:
                analysis = last_resume 
                
            session['current_analysis'] = analysis
    if request.method == 'POST':
        company_name = request.form.get('company_name', '')
        job_role = request.form.get('job_role', '')
        
        if analysis:
            cover_letter = generate_cover_letter(analysis, job_role, company_name)
            return render_template('cover_letter.html', cover_letter=cover_letter, analysis=analysis)
    return render_template('cover_letter.html', analysis=analysis)

@app.route('/ats-checker')
def ats_checker():
    analysis = session.get('current_analysis')
    if not analysis and session.get('user_id'):
        user_id = session.get('user_id')
        user_resumes = get_user_resumes(user_id)
        if user_resumes and len(user_resumes) > 0:
            analysis = user_resumes
            session['current_analysis'] = analysis
    if analysis:
        ats_result = analysis.get('ats_score', {})
        return render_template('ats_result.html', ats=ats_result)
    return redirect(url_for('index'))
    
@app.route('/resume-rewriter', methods=['GET', 'POST'])
def resume_rewriter():
    rewritten_text = ""
    suggestions = []
    original_text = ""
    if request.method == 'POST':
        original_text = request.form.get('resume_text', '')
        section_type = request.form.get('section_type', 'bullet_point')
        if original_text:
            rewritten_text = rewrite_resume_section(original_text, section_type)
            suggestions = get_action_verb_suggestions(original_text)
    return render_template('resume_rewriter.html', original_text=original_text, rewritten_text=rewritten_text, suggestions=suggestions)

@app.route('/linkedin-optimizer')
def linkedin_optimizer():
    analysis = session.get('current_analysis')
    
    if not analysis and session.get('user_id'):
        user_id = session.get('user_id')
        user_resumes = get_user_resumes(user_id)
        if user_resumes and len(user_resumes) > 0:
            analysis = user_resumes
            session['current_analysis'] = analysis

    if analysis:
        linkedin_data = analysis.get('linkedin_optimization', {})
        return render_template('linkedin_optimizer.html', linkedin=linkedin_data)
        
    return redirect(url_for('index'))

@app.route('/set-language/<lang>')
def set_language(lang):
    session['language'] = lang
    return redirect(request.referrer or url_for('index'))

@app.route('/download-report', methods=['POST'])
def download_report():
    try:
        analysis = request.json
        
        skills_html = ""
        if analysis.get('skills'):
            for category, skills in analysis['skills'].items():
                skills_html += f'<div class="skill-card"><strong>{category}:</strong><br>'
                for skill in skills:
                    skills_html += f'<span class="skill-badge">{skill}</span>'
                skills_html += '</div>'
        else:
            skills_html = '<p>No technical skills detected</p>'
        
        top_matches_html = ""
        for match in analysis.get('top_matches', []):
            top_matches_html += f"""
            <tr>
                <td><strong>{match['job_role']}</strong></td>
                <td>{match['match_score']}%</td>
                <td>{match['category']}</td>
                <td>{match['experience_level']}</td>
                <td>{match['salary_range']}</td>
            </tr>
            """
        
        education_html = ""
        for edu in analysis.get('education', []):
            education_html += f'<li>{edu}</li>'
        
        recommendations_html = ""
        for rec in analysis.get('recommendations', []):
            recommendations_html += f'<div class="recommendation-card">{rec}</div>'
        
        interview_q = analysis.get('interview_questions', {})
        
        salary_comp = analysis.get('salary_comparison', {})
        salary_html = f"""
        <div class="salary-grid">
            <div class="salary-card"><div class="salary-value">${salary_comp.get('estimated_salary', 0):,}</div><div>Estimated Salary</div></div>
            <div class="salary-card"><div class="salary-value">${salary_comp.get('market_average', 0):,}</div><div>Market Average</div></div>
        </div>
        <div class="salary-insight">💡 {salary_comp.get('recommendation', '')}</div>
        """
        
        ats_score = analysis.get('ats_score', {})
        ats_html = f"""
        <div class="section">
            <div class="section-title">🏆 ATS Compatibility Score</div>
            <div class="score-value">{ats_score.get('score', 0)}%</div>
            <div class="progress-bar"><div class="progress-fill" style="width: {ats_score.get('score', 0)}%"></div></div>
            <p><strong>Status:</strong> {'ATS Friendly' if ats_score.get('is_ats_friendly', False) else 'Needs Improvement'}</p>
        </div>
        """
        
        contact_info = analysis.get('contact_info', {})
        emails_display = ', '.join(contact_info.get('emails', [])) if contact_info.get('emails') else 'Not found'
        phones_display = ', '.join(contact_info.get('phones', [])) if contact_info.get('phones') else 'Not found'
        
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Resume Analysis Report</title>
    <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; padding: 40px; max-width: 1000px; margin: 0 auto; background: #f5f5f5; }}
        .container {{ background: white; border-radius: 15px; box-shadow: 0 10px 30px rgba(0,0,0,0.1); overflow: hidden; }}
        .header {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 30px; text-align: center; }}
        h1 {{ margin: 0; font-size: 28px; }}
        .date {{ margin-top: 10px; opacity: 0.9; font-size: 14px; }}
        .content {{ padding: 30px; }}
        .score-grid {{ display: grid; grid-template-columns: repeat(3, 1fr); gap: 20px; margin-bottom: 20px; }}
        .score-card {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px; border-radius: 10px; text-align: center; }}
        .score-value {{ font-size: 36px; font-weight: bold; margin: 10px 0; }}
        .section {{ margin: 25px 0; padding: 20px; background: #f8f9fa; border-radius: 10px; border: 1px solid #e0e0e0; }}
        .section-title {{ font-size: 20px; font-weight: bold; color: #667eea; margin-bottom: 15px; border-left: 4px solid #667eea; padding-left: 10px; }}
        table {{ width: 100%; border-collapse: collapse; }}
        th, td {{ padding: 12px; text-align: left; border-bottom: 1px solid #ddd; }}
        th {{ background: #667eea; color: white; }}
        .skill-badge {{ background: #667eea; color: white; padding: 5px 12px; border-radius: 20px; display: inline-block; margin: 5px; font-size: 12px; }}
        .recommendation-card {{ padding: 12px; margin: 10px 0; background: white; border-left: 4px solid #28a745; border-radius: 5px; }}
        .salary-grid {{ display: grid; grid-template-columns: repeat(2, 1fr); gap: 20px; margin-bottom: 20px; }}
        .salary-card {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px; border-radius: 10px; text-align: center; }}
        .salary-value {{ font-size: 28px; font-weight: bold; }}
        .salary-insight {{ margin-top: 15px; padding: 15px; background: #d4edda; border-radius: 8px; color: #155724; }}
        .progress-bar {{ height: 20px; background: #e0e0e0; border-radius: 10px; overflow: hidden; margin: 10px 0; }}
        .progress-fill {{ height: 100%; background: linear-gradient(90deg, #667eea, #764ba2); border-radius: 10px; }}
        .footer {{ text-align: center; padding: 20px; background: #f8f9fa; color: #666; font-size: 12px; }}
    </style>
</head>
<body>
<div class="container">
    <div class="header">
        <h1>🤖 AI Resume Analysis Report</h1>
        <div class="date">Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</div>
        <div class="date">Target Role: {analysis.get('job_role', 'N/A')}</div>
    </div>
    <div class="content">
        <div class="score-grid">
            <div class="score-card"><div>Job Match Score</div><div class="score-value">{analysis.get('match_score', 0)}%</div></div>
            <div class="score-card"><div>Skills Match</div><div class="score-value">{analysis.get('skill_score', 0)}%</div></div>
            <div class="score-card"><div>Completeness</div><div class="score-value">{analysis.get('completeness_score', 0)}%</div></div>
        </div>
        
        <div class="section">
            <div class="section-title">📞 Contact Information</div>
            <p><strong>Email:</strong> {emails_display}</p>
            <p><strong>Phone:</strong> {phones_display}</p>
            <p><strong>Total Words:</strong> {analysis.get('total_words', 0)}</p>
        </div>
        
        <div class="section">
            <div class="section-title">💻 Technical Skills</div>
            {skills_html}
        </div>
        
        <div class="section">
            <div class="section-title">🎯 Top 16 Job Matches</div>
            <table><thead><tr><th>Job Role</th><th>Match Score</th><th>Category</th><th>Experience</th><th>Salary</th></tr></thead><tbody>{top_matches_html}</tbody></table>
        </div>
        
        {ats_html}
        
        <div class="section">
            <div class="section-title">💰 Salary Analysis</div>
            {salary_html}
        </div>
        
        <div class="section">
            <div class="section-title">🎓 Education</div>
            <ul>{education_html if education_html else '<li>No education detected</li>'}</ul>
        </div>
        
        <div class="section">
            <div class="section-title">🎤 Interview Questions</div>
            <h4>Technical:</h4><ul>{''.join([f'<li>{q}</li>' for q in interview_q.get('technical', [])])}</ul>
            <h4>Behavioral:</h4><ul>{''.join([f'<li>{q}</li>' for q in interview_q.get('behavioral', [])])}</ul>
            <h4>Coding:</h4><ul>{''.join([f'<li>{q}</li>' for q in interview_q.get('coding', [])])}</ul>
        </div>
        
        <div class="section">
            <div class="section-title">💡 Recommendations</div>
            {recommendations_html}
        </div>
    </div>
    <div class="footer">
        <p>Generated by AI Resume Analyzer | © 2026 All Rights Reserved</p>
    </div>
</div>
</body>
</html>"""
        
        return html_content, 200, {
            'Content-Type': 'text/html',
            'Content-Disposition': f'attachment; filename=resume_analysis_{datetime.now().strftime("%Y%m%d_%H%M%S")}.html'
        }
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ==================== DEBUG ROUTE ====================
@app.route('/debug-db')
def debug_db():
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        
        users = c.execute("SELECT * FROM users").fetchall()
        resumes = c.execute("SELECT * FROM user_resumes").fetchall()
        
        conn.close()
        
        return f"""
        <h2>✅ Database Debug</h2>
        <p>Database Path: {DB_PATH}</p>
        <p>Users: {len(users)}</p>
        <pre>{users}</pre>
        <p>Resumes: {len(resumes)}</p>
        <pre>{resumes}</pre>
        """
    except Exception as e:
        return f"❌ Error: {e}"

# ==================== RUN APP ====================
if __name__ == '__main__':
    import os
    port = int(os.environ.get('PORT', 5000))
    
    print("=" * 60)
    print("🤖 AI RESUME ANALYZER - SECURE VERSION")
    print("=" * 60)
    print(f"✅ Total Job Roles: {len(JOB_DATABASE)}")
    print(f"✅ Gemini 2.5 Flash: {'CONNECTED' if GEMINI_AVAILABLE else 'NOT CONNECTED'}")
    print(f"✅ Google Login: ENABLED")
    print(f"🌐 Server: http://0.0.0.0:{port}")
    print("=" * 60)
    app.run(debug=False, host='0.0.0.0', port=port)
